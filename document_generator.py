from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from config import FormType, Config
import os
from enum import Enum
from abc import ABC, abstractmethod
import markdown
import jinja2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

class OutputFormat(Enum):
    DOCX = "docx"
    PDF = "pdf"
    MARKDOWN = "markdown"
    HTML = "html"

class BaseDocumentGenerator(ABC):
    def __init__(self, form_data: dict):
        self.form_data = form_data
        self.year_words = {
            "2024": "twenty four",
            "2025": "twenty five",
            "2026": "twenty six",
            "2027": "twenty seven",
            "2028": "twenty eight",
            "2029": "twenty nine",
            "2030": "thirty"
        }
        self.current_date = datetime.now()
        self.year_in_words = self.year_words.get(str(self.current_date.year), "")
    
    @abstractmethod
    def generate(self):
        pass
    
    @abstractmethod
    def save(self, output_path: str):
        pass

class DocxGenerator(BaseDocumentGenerator):
    def __init__(self, form_data: dict):
        super().__init__(form_data)
        self.document = Document()
        self._setup_document()
    
    def _setup_document(self):
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def _add_heading(self, text: str, level: int = 1, alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER):
        heading = self.document.add_heading('', level=level)
        heading.alignment = alignment
        run = heading.add_run(text)
        run.font.size = Pt(12)
        return heading
    
    def _add_paragraph(self, text: str, alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY):
        paragraph = self.document.add_paragraph(text)
        paragraph.alignment = alignment
        for run in paragraph.runs:
            run.font.size = Pt(12)
        return paragraph

    def generate(self) -> str:
        # Title section
        self._add_paragraph("R e s i d e n t i a l").alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_paragraph("Pakistan Defence Officers Housing Authority").alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_paragraph("K a r a c h i").alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_paragraph("Licence – 'a'").alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add divider lines
        self._add_paragraph("-" * 75).alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_paragraph("")
        self._add_paragraph("-" * 75).alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_paragraph("")
        self._add_paragraph("")
        
        # Repeat header
        self._add_paragraph("Pakistan Defence Officers Housing Authority").alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_paragraph("K a r a c h i").alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_paragraph("Licence 'a'").alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add main content
        self._add_paragraph(f"""THIS INDENTURE made this {self.current_date.strftime('%d')}th day of {self.current_date.strftime('%B')} (in the year two thousand and {self.year_in_words}) BETWEEN the Pakistan Defence Officers Housing Authority established under Article 4 of Pakistan Defence Officers Housing Authority Order, 1980, having its office at Korangi Road, Karachi (hereinafter called the "1st Party") AND {self.form_data['licensee_name']}\n{self.form_data['licensee_address']}\n(hereinafter called the "Licensee-2nd Party". The terms 1st Party and 2nd Party shall include their respective executors, successors-in-interest and assigns).""")
        
        # Add KPT details
        self._add_paragraph(f"""WHEREAS the (KPT) Karachi Port Trust (hereinafter referred to as the lessor) through a deed registered in the office of the Sub-Registrar {self.form_data['sub_registrar']} Town, Karachi, as No {self.form_data['kpt_book_no']} Book-I dated {self.form_data['kpt_book_date'].strftime('%d-%m-%Y')}, M.F. Roll No.{self.form_data['kpt_mf_roll_no']} dated {self.form_data['kpt_mf_roll_date'].strftime('%d-%m-%Y')} admeasuring ­­­­­­­{self.form_data['land_size']} acres, had authorised the Pakistan Defence Officers Housing Authority, Karachi (hereinafter called the Authority) to enter upon the entire area of land shown in the plan attached to the lease including the plot referred to hereinafter for the purpose of developing it and for the construction of building, possession whereof had already been taken over by the Authority subject to, the terms and conditions contained in the Agreement;""")
        
        # Add possession clause
        self._add_paragraph(f"""AND WHEREAS the 1st Party now being fully entitled to seize and well possessed of all the piece and parcel of land measuring {self.form_data['land_size']} acres of land in Deh {self.form_data['deh']} bearing survey sheet No. {self.form_data['survey_sheet_number']} and fully described in the Schedule hereunder and fully competent and legally entitled as owners to allot the same.""")
        
        # Add transfer clause
        self._add_paragraph(f"""AND WHEREAS the licensee has been allotted / transferred vide allotment / transfer order No {self.form_data['transfer_order_no']} dated {self.form_data['transfer_order_date'].strftime('%d-%m-%Y')} the plot bearing No. {self.form_data['plot_number']} Survey Sheet No.{self.form_data['survey_sheet_number']} in the territorial division of {self.form_data['territorial_division']} Police Station in the layout plan of the entire area measuring {self.form_data['land_size']} acres as shown in the Schedule hereunder.""")
        
        # Add page number
        self._add_paragraph("2").alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add license clause
        self._add_paragraph("AND WHEREAS under the licence the 2nd Party is entitled to enter upon the said plot of land for the purpose of constructing a building thereon.")
        
        # Add witnesseth clause
        self._add_paragraph("NOW THIS INDENTURE WITNESSETH as follows :--")
        
        # Add numbered clauses
        self._add_paragraph("1.\tThat the 1st Party do hereby authorise and permit the 2nd Party to enter upon the said plot of land for the purpose of constructing a building thereon in accordance with the terms and conditions hereinafter following :--")
        
        # Add subclauses
        subclauses = [
            "(i)\tThe 2nd Party shall at his own cost and within the period of 2 years from the date of execution of this licence erect, complete and finish upon the said plot a residential house in accordance with the plan and design approved by the competent authority (hereinafter called the \"Authority\") subject to the condition that no construction work shall be started by the 2nd Party on his plot unless the preliminary stages of development shall have been completed and permission in writing shall have been obtained from the Authority. In the event of the licensee failing to comply with the conditions hereinafter appearing the 1st Party may at his discretion recover from the 2nd Party as agreed liquidated damages and not by way of penalty a sum equal to half per sent of the estimated cost of work remaining incomplete for every month the work remains incomplete subject to maximum of 5 per cent of the estimated cost of the work remaining un-finished after the due date. Provided that if the licensee fails to complete and finish the building by the date finally fixed by the Authority, the 1st Party may terminate this licence and resume the plot and any structures erected thereon.",
            "(ii)\tThe construction shall be done in accordance with the building bye-laws and the rules laid down by the Authority (1st Party).",
            "(iii)\tWith the execution of these presents the rights and liabilities accrued under this instrument shall devolve upon the 2nd Party and he shall be bound by such terms and conditions of the licence as are expressly or by necessary implication or analogy applicable to him.",
            "(iv)\tThis is a licence with permission to build and occupy. After the completion of the building a proper lease will be given to the Licensee for a period of 99 years by the (1st Party) on such terms and conditions as they deem necessary or may be imposed by the Government or any other Authority.",
            "(v)\tThe Licensee shall deposit with any scheduled bank duly authorised by the 1st Party or with the 1st Party :",
            f"(a)\tThe amount at the rate of Rs.{self.form_data['premium_rate']} per square yard to be paid in lump sum before execution of this licence agreement towards the premium of the plot.",
            f"(b)\tThe ground rent is payable in advance on or before the first day of July every year at the rate of {self.form_data['ground_rent_rate']} paisas per square yard per annum. The first payment shall be made on the first day of July, next following the day when the licensee takes possession of the plot allotted/transferred to him/her,",
            "(vi)\tThe 2nd Party shall pay all the calls (hereinafter called the \"development charges\") levied by the 1st Party from time to time at their office for an amount equal to the proportion of expenses to be incurred by the (1st Party) on the execution and completion of the development schemes. The decision of the Executive Board of the 1st Party as to the amount so payable shall be final and binding on the licensee."
        ]
        
        for subclause in subclauses:
            self._add_paragraph(subclause)
        
        # Add page 3
        self._add_paragraph("3").alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Additional clauses
        self._add_paragraph("(vii)\tAll arrears of payments due and payable by the Licensee shall be recoverable as arrears of land revenues.")
        
        self._add_paragraph("2.\tIt is hereby agreed that on the completion of the building in accordance with the said terms and conditions and on the licensee complying with the said rules he shall be entitled to a lease of the said plot for 99 years in the form prescribed by the Executive Board of the 1st Party and IT IS HEREBY FURTHER AGREED that until such lease has been granted by the (1st Party) the licensee shall not have any right or interest in the said plot except that of a bare licensee and shall not without the previous permission in writing of the (1st Party) transfer his interest in the area allotted to him either in part or whole except for the purpose of raising loans from the House Building Finance Corporation, authorised banks and insurance companies for construction of building thereon.")
        
        self._add_paragraph("3.\tShould the licensee commit breach of any of the terms and conditions of these presents or should he neglect to comply with any direction given to him by the 1st Party or in any other respect fail to carry out his obligations under these presents for reasons not beyond his control or fail to pay development charges or other dues, the 1st Party shall have the right to terminate this licence and on such termination the payment made by him to the 1st Party shall be forfeited and he shall have no further claim whatsoever against the 1st Party except at the option of 1st Party he may receive compensation to the extent of the amount of actual expenditure incurred by him on the plot.")
        
        self._add_paragraph("Provided that the 1st Party may in his absolute discretion have the building sold out either by public auction or private contract, in which case the licensee shall be entitled to the net sale proceeds of the building or to the amount of actual expenditure incurred by him on having the building constructed whichever is less.")
        
        # Add schedule
        schedule_heading = "THE SCHEDULE ABOVE REFERRED TO"
        self._add_paragraph(schedule_heading).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        schedule = f"""ALL that piece and parcel of land measuring {self.form_data['plot_area']} square yards bearing Plot No {self.form_data['plot_number']} and bounded
North................... {self.form_data['north_boundary']}
South ................{self.form_data['south_boundary']}
East................... {self.form_data['east_boundary']}
West................... {self.form_data['west_boundary']}
Situated in Police Station {self.form_data['police_station']}"""
        
        self._add_paragraph(schedule)
        
        # Add page 4
        self._add_paragraph("4").alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add empty lines
        for _ in range(3):
            self._add_paragraph("")
        
        # Add signature section
        sec_sign = """SIGNED by the Secretary, Pakistan\t\t\t\t    …………………………………………...
Defence Officers Housing Authority Karachi.\t\t\t               Signature of the Secretary
In the presence of :\t\t\t\t\t\t                 1st Party"""
        
        self._add_paragraph(sec_sign)
        
        self._add_paragraph("1.")
        self._add_paragraph("")
        self._add_paragraph("")
        self._add_paragraph("2.")
        self._add_paragraph("")
        self._add_paragraph("")
        self._add_paragraph("\t\t\t\t\t\t    …………………………………………...\n\t\t\t\t\t\t\t\t Member Executive Board")
        
        for _ in range(5):
            self._add_paragraph("")
        
        # Add witness section
        witness_sec = f"""SIGNED by the above named\t\t\t\t    …………………………………………...
Licensee – 2nd Party in the\t\t\t\t\t        Licensee / 2nd Party
Presence of :
Witness: (1)  Signature:……………………………..
\tName: {self.form_data['witness1_name']}
\tAddress: {self.form_data['witness1_address']}
\tCNIC #: {self.form_data['witness1_cnic']}


  (2)  Signature:……………………………..
\tName: {self.form_data['witness2_name']}
\tAddress: {self.form_data['witness2_address']}
\tCNIC #: {self.form_data['witness2_cnic']}"""
        
        self._add_paragraph(witness_sec)
        
        return ""
    
    def save(self, output_path: str):
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.document.save(output_path)
        
        
class MarkdownGenerator(BaseDocumentGenerator):
    def generate(self) -> str:
        template = """
# R e s i d e n t i a l

# Pakistan Defence Officers Housing Authority
# K a r a c h i

# Licence – 'a'




-------------------


-------------------








# Pakistan Defence Officers Housing Authority
# K a r a c h i

# Licence 'a'

THIS INDENTURE made this {{current_date}} BETWEEN the Pakistan Defence Officers Housing Authority established under Article 4 of Pakistan Defence Officers Housing Authority Order, 1980, having its office at Korangi Road, Karachi (hereinafter called the "1st Party") AND {{licensee_name}}  
{{licensee_address}}  
(hereinafter called the "Licensee-2nd Party". The terms 1st Party and 2nd Party shall include their respective executors, successors-in-interest and assigns).

WHEREAS the (KPT) Karachi Port Trust (hereinafter referred to as the lessor) through a deed registered in the office of the Sub-Registrar {{sub_registrar}} Town, Karachi, as No {{kpt_book_no}} Book-I dated {{kpt_book_date}}, M.F. Roll No.{{kpt_mf_roll_no}} dated {{kpt_mf_roll_date}} admeasuring {{land_size}} acres, had authorised the Pakistan Defence Officers Housing Authority, Karachi (hereinafter called the Authority) to enter upon the entire area of land shown in the plan attached to the lease including the plot referred to hereinafter for the purpose of developing it and for the construction of building, possession whereof had already been taken over by the Authority subject to, the terms and conditions contained in the Agreement;

AND WHEREAS the 1st Party now being fully entitled to seize and well possessed of all the piece and parcel of land measuring {{land_size}} acres of land in Deh {{deh}} bearing survey sheet No. {{survey_sheet_number}} and fully described in the Schedule hereunder and fully competent and legally entitled as owners to allot the same.

AND WHEREAS the licensee has been allotted / transferred vide allotment / transfer order No {{transfer_order_no}} dated {{transfer_order_date}} the plot bearing No. {{plot_number}} Survey Sheet No.{{survey_sheet_number}} in the territorial division of {{territorial_division}} Police Station in the layout plan of the entire area measuring {{land_size}} acres as shown in the Schedule hereunder.

# 2

AND WHEREAS under the licence the 2nd Party is entitled to enter upon the said plot of land for the purpose of constructing a building thereon.

NOW THIS INDENTURE WITNESSETH as follows :--

1. That the 1st Party do hereby authorise and permit the 2nd Party to enter upon the said plot of land for the purpose of constructing a building thereon in accordance with the terms and conditions hereinafter following :--

    (i) The 2nd Party shall at his own cost and within the period of 2 years from the date of execution of this licence erect, complete and finish upon the said plot a residential house in accordance with the plan and design approved by the competent authority (hereinafter called the "Authority") subject to the condition that no construction work shall be started by the 2nd Party on his plot unless the preliminary stages of development shall have been completed and permission in writing shall have been obtained from the Authority. In the event of the licensee failing to comply with the conditions hereinafter appearing the 1st Party may at his discretion recover from the 2nd Party as agreed liquidated damages and not by way of penalty a sum equal to half per sent of the estimated cost of work remaining incomplete for every month the work remains incomplete subject to maximum of 5 per cent of the estimated cost of the work remaining un-finished after the due date. Provided that if the licensee fails to complete and finish the building by the date finally fixed by the Authority, the 1st Party may terminate this licence and resume the plot and any structures erected thereon.

    (ii) The construction shall be done in accordance with the building bye-laws and the rules laid down by the Authority (1st Party).

    (iii) With the execution of these presents the rights and liabilities accrued under this instrument shall devolve upon the 2nd Party and he shall be bound by such terms and conditions of the licence as are expressly or by necessary implication or analogy applicable to him.

    (iv) This is a licence with permission to build and occupy. After the completion of the building a proper lease will be given to the Licensee for a period of 99 years by the (1st Party) on such terms and conditions as they deem necessary or may be imposed by the Government or any other Authority.

    (v) The Licensee shall deposit with any scheduled bank duly authorised by the 1st Party or with the 1st Party :

    (a) The amount at the rate of Rs.{{premium_rate}} per square yard to be paid in lump sum before execution of this licence agreement towards the premium of the plot.

    (b) The ground rent is payable in advance on or before the first day of July every year at the rate of {{ground_rent_rate}} paisas per square yard per annum. The first payment shall be made on the first day of July, next following the day when the licensee takes possession of the plot allotted/transferred to him/her,

    (vi) The 2nd Party shall pay all the calls (hereinafter called the "development charges") levied by the 1st Party from time to time at their office for an amount equal to the proportion of expenses to be incurred by the (1st Party) on the execution and completion of the development schemes. The decision of the Executive Board of the 1st Party as to the amount so payable shall be final and binding on the licensee.

# 3

(vii) All arrears of payments due and payable by the Licensee shall be recoverable as arrears of land revenues.

2. It is hereby agreed that on the completion of the building in accordance with the said terms and conditions and on the licensee complying with the said rules he shall be entitled to a lease of the said plot for 99 years in the form prescribed by the Executive Board of the 1st Party and IT IS HEREBY FURTHER AGREED that until such lease has been granted by the (1st Party) the licensee shall not have any right or interest in the said plot except that of a bare licensee and shall not without the previous permission in writing of the (1st Party) transfer his interest in the area allotted to him either in part or whole except for the purpose of raising loans from the House Building Finance Corporation, authorised banks and insurance companies for construction of building thereon.

3. Should the licensee commit breach of any of the terms and conditions of these presents or should he neglect to comply with any direction given to him by the 1st Party or in any other respect fail to carry out his obligations under these presents for reasons not beyond his control or fail to pay development charges or other dues, the 1st Party shall have the right to terminate this licence and on such termination the payment made by him to the 1st Party shall be forfeited and he shall have no further claim whatsoever against the 1st Party except at the option of 1st Party he may receive compensation to the extent of the amount of actual expenditure incurred by him on the plot.

Provided that the 1st Party may in his absolute discretion have the building sold out either by public auction or private contract, in which case the licensee shall be entitled to the net sale proceeds of the building or to the amount of actual expenditure incurred by him on having the building constructed whichever is less.


# THE SCHEDULE ABOVE REFERRED TO


ALL that piece and parcel of land measuring {{plot_area}} square yards bearing Plot No {{plot_number}} and bounded  
North................... {{north_boundary}}  
South ................{{south_boundary}}  
East................... {{east_boundary}}  
West................... {{west_boundary}}  
Situated in Police Station {{police_station}}

# 4






SIGNED by the Secretary, Pakistan                    …………………………………………...  
Defence Officers Housing Authority Karachi.                      Signature of the Secretary  
In the presence of :                                             1st Party  

1.



2.


                                            …………………………………………...  
                                                     Member Executive Board  











SIGNED by the above named                            …………………………………………...  
Licensee – 2nd Party in the                                 Licensee / 2nd Party  
Presence of :  
Witness: (1)  Signature:……………………………..  
        Name: {{witness1_name}}  
        Address: {{witness1_address}}  
        CNIC #: {{witness1_cnic}}  


  (2)  Signature:……………………………..  
        Name: {{witness2_name}}  
        Address: {{witness2_address}}  
        CNIC #: {{witness2_cnic}}  
"""
        return jinja2.Template(template).render(**self.form_data, current_date=self.current_date.strftime('%d %B %Y'))
    
    def save(self, output_path: str):
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, 'w') as f:
            f.write(self.generate())

class HtmlGenerator(BaseDocumentGenerator):
    def generate(self) -> str:
        template = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>DHA License Document</title>
    <style>
        @page {
            size: A4;
            margin: 2.5cm;
        }
        body {
            font-family: Arial, sans-serif;
            line-height: 1.6;
            max-width: 21cm;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }
        .title {
            text-align: center;
            font-size: 16px;
            letter-spacing: 2px;
            margin-bottom: 10px;
        }
        .subtitle {
            text-align: center;
            font-size: 14px;
            margin-bottom: 8px;
        }
        .divider {
            border: none;
            margin: 20px 0;
            width: 100%;
            text-align: center;
        }
        .content {
            text-align: justify;
            font-size: 12px;
            margin: 20px 0;
            text-indent: 36px;
        }
        .clause {
            text-align: justify;
            font-size: 12px;
            margin: 20px 0;
            padding-left: 36px;
            text-indent: -36px;
        }
        .subclause {
            text-align: justify;
            font-size: 12px;
            margin: 10px 0;
            padding-left: 72px;
            text-indent: -36px;
        }
        .schedule {
            margin: 30px 0;
        }
        .schedule-content {
            text-align: left;
            font-size: 12px;
            margin: 20px 0;
            white-space: pre-wrap;
        }
        .page-number {
            font-size: 12px;
            margin-left: 36px;
            margin-top: 20px;
            margin-bottom: 20px;
        }
        .signature-section {
            margin-top: 50px;
            white-space: pre-wrap;
            font-size: 12px;
        }
        .witness-section {
            margin-top: 30px;
            white-space: pre-wrap;
            font-size: 12px;
        }
        .page-break {
            page-break-before: always;
        }
        .spacer {
            height: 48px;
        }
        .big-spacer {
            height: 72px;
        }
        @media print {
            body { font-size: 12pt; }
            .page-break { page-break-before: always; }
        }
    </style>
</head>
<body>
    <div class="title">R e s i d e n t i a l</div>
    <div class="spacer"></div>
    <div class="title">Pakistan Defence Officers Housing Authority</div>
    <div class="title">K a r a c h i</div>
    <div class="title">Licence – 'a'</div>
    <div class="spacer"></div>
    <div class="divider">---------------------------------------------------------------------------</div>
    <div class="spacer"></div>
    <div class="divider">---------------------------------------------------------------------------</div>
    <div class="big-spacer"></div>
    <div class="title">Pakistan Defence Officers Housing Authority</div>
    <div class="title">K a r a c h i</div>
    <div class="title">Licence 'a'</div>
    
    <div class="content">
        THIS INDENTURE made this {{current_date}} BETWEEN the Pakistan Defence Officers Housing Authority established under Article 4 of Pakistan Defence Officers Housing Authority Order, 1980, having its office at Korangi Road, Karachi (hereinafter called the "1st Party") AND {{licensee_name}}<br>
        {{licensee_address}}<br>
        (hereinafter called the "Licensee-2nd Party". The terms 1st Party and 2nd Party shall include their respective executors, successors-in-interest and assigns).
    </div>
    
    <div class="content">
        WHEREAS the (KPT) Karachi Port Trust (hereinafter referred to as the lessor) through a deed registered in the office of the Sub-Registrar {{sub_registrar}} Town, Karachi, as No {{kpt_book_no}} Book-I dated {{kpt_book_date}}, M.F. Roll No.{{kpt_mf_roll_no}} dated {{kpt_mf_roll_date}} admeasuring {{land_size}} acres, had authorised the Pakistan Defence Officers Housing Authority, Karachi (hereinafter called the Authority) to enter upon the entire area of land shown in the plan attached to the lease including the plot referred to hereinafter for the purpose of developing it and for the construction of building, possession whereof had already been taken over by the Authority subject to, the terms and conditions contained in the Agreement;
    </div>

    <div class="content">
        AND WHEREAS the 1st Party now being fully entitled to seize and well possessed of all the piece and parcel of land measuring {{land_size}} acres of land in Deh {{deh}} bearing survey sheet No. {{survey_sheet_number}} and fully described in the Schedule hereunder and fully competent and legally entitled as owners to allot the same.
    </div>

    <div class="content">
        AND WHEREAS the licensee has been allotted / transferred vide allotment / transfer order No {{transfer_order_no}} dated {{transfer_order_date}} the plot bearing No. {{plot_number}} Survey Sheet No.{{survey_sheet_number}} in the territorial division of {{territorial_division}} Police Station in the layout plan of the entire area measuring {{land_size}} acres as shown in the Schedule hereunder.
    </div>

    <div class="page-break"></div>
    
    <div class="content">
        AND WHEREAS under the licence the 2nd Party is entitled to enter upon the said plot of land for the purpose of constructing a building thereon.
    </div>

    <div class="content">
        NOW THIS INDENTURE WITNESSETH as follows :--
    </div>

    <div class="clause">
        1. That the 1st Party do hereby authorise and permit the 2nd Party to enter upon the said plot of land for the purpose of constructing a building thereon in accordance with the terms and conditions hereinafter following :--
    </div>

    <div class="subclause">
        (i) The 2nd Party shall at his own cost and within the period of 2 years from the date of execution of this licence erect, complete and finish upon the said plot a residential house in accordance with the plan and design approved by the competent authority (hereinafter called the "Authority") subject to the condition that no construction work shall be started by the 2nd Party on his plot unless the preliminary stages of development shall have been completed and permission in writing shall have been obtained from the Authority. In the event of the licensee failing to comply with the conditions hereinafter appearing the 1st Party may at his discretion recover from the 2nd Party as agreed liquidated damages and not by way of penalty a sum equal to half per sent of the estimated cost of work remaining incomplete for every month the work remains incomplete subject to maximum of 5 per cent of the estimated cost of the work remaining un-finished after the due date. Provided that if the licensee fails to complete and finish the building by the date finally fixed by the Authority, the 1st Party may terminate this licence and resume the plot and any structures erected thereon.
    </div>

    <div class="subclause">
        (ii) The construction shall be done in accordance with the building bye-laws and the rules laid down by the Authority (1st Party).
    </div>

    <div class="subclause">
        (iii) With the execution of these presents the rights and liabilities accrued under this instrument shall devolve upon the 2nd Party and he shall be bound by such terms and conditions of the licence as are expressly or by necessary implication or analogy applicable to him.
    </div>

    <div class="subclause">
        (iv) This is a licence with permission to build and occupy. After the completion of the building a proper lease will be given to the Licensee for a period of 99 years by the (1st Party) on such terms and conditions as they deem necessary or may be imposed by the Government or any other Authority.
    </div>

    <div class="subclause">
        (v) The Licensee shall deposit with any scheduled bank duly authorised by the 1st Party or with the 1st Party :
    </div>

    <div class="subclause">
        (a) The amount at the rate of Rs.{{premium_rate}} per square yard to be paid in lump sum before execution of this licence agreement towards the premium of the plot.
    </div>

    <div class="subclause">
        (b) The ground rent is payable in advance on or before the first day of July every year at the rate of {{ground_rent_rate}} paisas per square yard per annum. The first payment shall be made on the first day of July, next following the day when the licensee takes possession of the plot allotted/transferred to him/her,
    </div>

    <div class="subclause">
        (vi) The 2nd Party shall pay all the calls (hereinafter called the "development charges") levied by the 1st Party from time to time at their office for an amount equal to the proportion of expenses to be incurred by the (1st Party) on the execution and completion of the development schemes. The decision of the Executive Board of the 1st Party as to the amount so payable shall be final and binding on the licensee.
    </div>

    <div class="page-break"></div>

    <div class="subclause">
        (vii) All arrears of payments due and payable by the Licensee shall be recoverable as arrears of land revenues.
    </div>

    <div class="clause">
        2. It is hereby agreed that on the completion of the building in accordance with the said terms and conditions and on the licensee complying with the said rules he shall be entitled to a lease of the said plot for 99 years in the form prescribed by the Executive Board of the 1st Party and IT IS HEREBY FURTHER AGREED that until such lease has been granted by the (1st Party) the licensee shall not have any right or interest in the said plot except that of a bare licensee and shall not without the previous permission in writing of the (1st Party) transfer his interest in the area allotted to him either in part or whole except for the purpose of raising loans from the House Building Finance Corporation, authorised banks and insurance companies for construction of building thereon.
    </div>

    <div class="clause">
        3. Should the licensee commit breach of any of the terms and conditions of these presents or should he neglect to comply with any direction given to him by the 1st Party or in any other respect fail to carry out his obligations under these presents for reasons not beyond his control or fail to pay development charges or other dues, the 1st Party shall have the right to terminate this licence and on such termination the payment made by him to the 1st Party shall be forfeited and he shall have no further claim whatsoever against the 1st Party except at the option of 1st Party he may receive compensation to the extent of the amount of actual expenditure incurred by him on the plot.
    </div>

    <div class="content">
        Provided that the 1st Party may in his absolute discretion have the building sold out either by public auction or private contract, in which case the licensee shall be entitled to the net sale proceeds of the building or to the amount of actual expenditure incurred by him on having the building constructed whichever is less.
    </div>

    <div class="title">THE SCHEDULE ABOVE REFERRED TO</div>

    <div class="schedule-content">
ALL that piece and parcel of land measuring {{plot_area}} square yards bearing Plot No {{plot_number}} and bounded
North................... {{north_boundary}}
South ................{{south_boundary}}
East................... {{east_boundary}}
West................... {{west_boundary}}
Situated in Police Station {{police_station}}
    </div>

    <div class="page-break"></div>
    <div class="big-spacer"></div>

    <div class="signature-section">
SIGNED by the Secretary, Pakistan                    …………………………………………...
Defence Officers Housing Authority Karachi.                      Signature of the Secretary
In the presence of :                                             1st Party

1.



2.


                                            …………………………………………...
                                                     Member Executive Board
    </div>

    <div class="big-spacer"></div>
    <div class="big-spacer"></div>
    <div class="witness-section">
        SIGNED by the above named                            …………………………………………...
        Licensee – 2nd Party in the                                 Licensee / 2nd Party
        Presence of :
        Witness: (1)  Signature:……………………………..
                      Name: {{witness1_name}}
                      Address: {{witness1_address}}
                      CNIC #: {{witness1_cnic}}


                 (2)  Signature:……………………………..
                      Name: {{witness2_name}}
                      Address: {{witness2_address}}
                      CNIC #: {{witness2_cnic}}
    </div>
</body>
</html>"""
        return jinja2.Template(template).render(**self.form_data, current_date=self.current_date.strftime('%d %B %Y'))
    
    def save(self, output_path: str):
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, 'w') as f:
            f.write(self.generate())

class PDFGenerator(BaseDocumentGenerator):
    def __init__(self, form_data: dict):
        super().__init__(form_data)
        self.styles = getSampleStyleSheet()
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup custom styles for the PDF document"""
        self.styles.add(ParagraphStyle(
            name='DHATitle',
            parent=self.styles['Title'],
            fontSize=14,
            spaceAfter=6,
            alignment=TA_CENTER,
            leading=16
        ))
        self.styles.add(ParagraphStyle(
            name='DHASubTitle',
            parent=self.styles['Title'],
            fontSize=12,
            spaceAfter=4,
            alignment=TA_CENTER,
            leading=14
        ))
        self.styles.add(ParagraphStyle(
            name='DHABody',
            parent=self.styles['Normal'],
            fontSize=11,
            leading=13,
            alignment=TA_JUSTIFY,
            firstLineIndent=36  # Add indentation for paragraphs
        ))
        self.styles.add(ParagraphStyle(
            name='DHAHeading',
            parent=self.styles['Heading1'],
            fontSize=12,
            spaceAfter=12,
            alignment=TA_CENTER,
            leading=14
        ))
        self.styles.add(ParagraphStyle(
            name='DHAPageNumber',
            parent=self.styles['Normal'],
            fontSize=11,
            alignment=TA_LEFT,
            leading=14,
            leftIndent=36
        ))
        self.styles.add(ParagraphStyle(
            name='DHAClause',
            parent=self.styles['Normal'],
            fontSize=11,
            leading=13,
            alignment=TA_JUSTIFY,
            leftIndent=36,
            firstLineIndent=-36
        ))

    def generate(self) -> list:
        story = []
        
        # Title section with exact spacing
        story.append(Spacer(1, 30))
        story.append(Paragraph("R e s i d e n t i a l", self.styles['DHATitle']))
        story.append(Spacer(1, 12))
        story.append(Paragraph("Pakistan Defence Officers Housing Authority", self.styles['DHATitle']))
        story.append(Paragraph("K a r a c h i", self.styles['DHATitle']))
        story.append(Spacer(1, 6))
        story.append(Paragraph("Licence – 'a'", self.styles['DHATitle']))
        story.append(Spacer(1, 36))
        
        # Add divider lines with exact spacing
        story.append(Paragraph("-" * 75, self.styles['DHAHeading']))
        story.append(Spacer(1, 36))
        story.append(Paragraph("-" * 75, self.styles['DHAHeading']))
        story.append(Spacer(1, 48))
        
        # Repeat header
        story.append(Paragraph("Pakistan Defence Officers Housing Authority", self.styles['DHATitle']))
        story.append(Paragraph("K a r a c h i", self.styles['DHATitle']))
        story.append(Spacer(1, 6))
        story.append(Paragraph("Licence 'a'", self.styles['DHATitle']))
        story.append(Spacer(1, 24))
        
        # Main content with proper indentation and spacing
        main_content = f"""THIS INDENTURE made this {self.current_date.strftime('%d')}th day of {self.current_date.strftime('%B')} (in the year two thousand and {self.year_in_words}) BETWEEN the Pakistan Defence Officers Housing Authority established under Article 4 of Pakistan Defence Officers Housing Authority Order, 1980, having its office at Korangi Road, Karachi (hereinafter called the "1st Party") AND {self.form_data['licensee_name']}\n{self.form_data['licensee_address']}\n(hereinafter called the "Licensee-2nd Party". The terms 1st Party and 2nd Party shall include their respective executors, successors-in-interest and assigns)."""
        
        story.append(Paragraph(main_content, self.styles['DHABody']))
        story.append(Spacer(1, 24))
        
        # Add KPT details
        kpt_details = f"""WHEREAS the (KPT) Karachi Port Trust (hereinafter referred to as the lessor) through a deed registered in the office of the Sub-Registrar {self.form_data['sub_registrar']} Town, Karachi, as No {self.form_data['kpt_book_no']} Book-I dated {self.form_data['kpt_book_date'].strftime('%d-%m-%Y')}, M.F. Roll No.{self.form_data['kpt_mf_roll_no']} dated {self.form_data['kpt_mf_roll_date'].strftime('%d-%m-%Y')} admeasuring ­­­­­­­{self.form_data['land_size']} acres, had authorised the Pakistan Defence Officers Housing Authority, Karachi (hereinafter called the Authority) to enter upon the entire area of land shown in the plan attached to the lease including the plot referred to hereinafter for the purpose of developing it and for the construction of building, possession whereof had already been taken over by the Authority subject to, the terms and conditions contained in the Agreement;"""
        
        story.append(Paragraph(kpt_details, self.styles['DHABody']))
        story.append(Spacer(1, 24))
        
        # Add possession clause
        possession_clause = f"""AND WHEREAS the 1st Party now being fully entitled to seize and well possessed of all the piece and parcel of land measuring {self.form_data['land_size']} acres of land in Deh {self.form_data['deh']} bearing survey sheet No. {self.form_data['survey_sheet_number']} and fully described in the Schedule hereunder and fully competent and legally entitled as owners to allot the same."""
        
        story.append(Paragraph(possession_clause, self.styles['DHABody']))
        story.append(Spacer(1, 24))
        
        # Add transfer clause
        transfer_clause = f"""AND WHEREAS the licensee has been allotted / transferred vide allotment / transfer order No {self.form_data['transfer_order_no']} dated {self.form_data['transfer_order_date'].strftime('%d-%m-%Y')} the plot bearing No. {self.form_data['plot_number']} Survey Sheet No.{self.form_data['survey_sheet_number']} in the territorial division of {self.form_data['territorial_division']} Police Station in the layout plan of the entire area measuring {self.form_data['land_size']} acres as shown in the Schedule hereunder."""
        
        story.append(Paragraph(transfer_clause, self.styles['DHABody']))
        
        # Add page number
        story.append(PageBreak())
        story.append(Paragraph("2", self.styles['DHAPageNumber']))
        story.append(Spacer(1, 12))
        
        # Add license clause
        story.append(Paragraph("AND WHEREAS under the licence the 2nd Party is entitled to enter upon the said plot of land for the purpose of constructing a building thereon.", self.styles['DHABody']))
        story.append(Spacer(1, 24))
        
        # Add witnesseth clause
        story.append(Paragraph("NOW THIS INDENTURE WITNESSETH as follows :--", self.styles['DHABody']))
        story.append(Spacer(1, 24))
        
        # Add numbered clauses with proper indentation
        story.append(Paragraph("1.\tThat the 1st Party do hereby authorise and permit the 2nd Party to enter upon the said plot of land for the purpose of constructing a building thereon in accordance with the terms and conditions hereinafter following :--", self.styles['DHAClause']))
        story.append(Spacer(1, 12))
        
        # Add subclauses with proper indentation
        subclauses = [
            "(i)\tThe 2nd Party shall at his own cost and within the period of 2 years from the date of execution of this licence erect, complete and finish upon the said plot a residential house in accordance with the plan and design approved by the competent authority (hereinafter called the \"Authority\") subject to the condition that no construction work shall be started by the 2nd Party on his plot unless the preliminary stages of development shall have been completed and permission in writing shall have been obtained from the Authority. In the event of the licensee failing to comply with the conditions hereinafter appearing the 1st Party may at his discretion recover from the 2nd Party as agreed liquidated damages and not by way of penalty a sum equal to half per sent of the estimated cost of work remaining incomplete for every month the work remains incomplete subject to maximum of 5 per cent of the estimated cost of the work remaining un-finished after the due date. Provided that if the licensee fails to complete and finish the building by the date finally fixed by the Authority, the 1st Party may terminate this licence and resume the plot and any structures erected thereon.",
            "(ii)\tThe construction shall be done in accordance with the building bye-laws and the rules laid down by the Authority (1st Party).",
            "(iii)\tWith the execution of these presents the rights and liabilities accrued under this instrument shall devolve upon the 2nd Party and he shall be bound by such terms and conditions of the licence as are expressly or by necessary implication or analogy applicable to him.",
            "(iv)\tThis is a licence with permission to build and occupy. After the completion of the building a proper lease will be given to the Licensee for a period of 99 years by the (1st Party) on such terms and conditions as they deem necessary or may be imposed by the Government or any other Authority.",
            "(v)\tThe Licensee shall deposit with any scheduled bank duly authorised by the 1st Party or with the 1st Party :",
            f"(a)\tThe amount at the rate of Rs.{self.form_data['premium_rate']} per square yard to be paid in lump sum before execution of this licence agreement towards the premium of the plot.",
            f"(b)\tThe ground rent is payable in advance on or before the first day of July every year at the rate of {self.form_data['ground_rent_rate']} paisas per square yard per annum. The first payment shall be made on the first day of July, next following the day when the licensee takes possession of the plot allotted/transferred to him/her,",
            "(vi)\tThe 2nd Party shall pay all the calls (hereinafter called the \"development charges\") levied by the 1st Party from time to time at their office for an amount equal to the proportion of expenses to be incurred by the (1st Party) on the execution and completion of the development schemes. The decision of the Executive Board of the 1st Party as to the amount so payable shall be final and binding on the licensee."
        ]
        
        for subclause in subclauses:
            story.append(Paragraph(subclause, self.styles['DHAClause']))
            story.append(Spacer(1, 6))
        
        # Add page 3
        story.append(PageBreak())
        story.append(Paragraph("3", self.styles['DHAPageNumber']))
        story.append(Spacer(1, 12))
        
        # Additional clauses
        story.append(Paragraph("(vii)\tAll arrears of payments due and payable by the Licensee shall be recoverable as arrears of land revenues.", self.styles['DHAClause']))
        story.append(Spacer(1, 24))
        
        story.append(Paragraph("2.\tIt is hereby agreed that on the completion of the building in accordance with the said terms and conditions and on the licensee complying with the said rules he shall be entitled to a lease of the said plot for 99 years in the form prescribed by the Executive Board of the 1st Party and IT IS HEREBY FURTHER AGREED that until such lease has been granted by the (1st Party) the licensee shall not have any right or interest in the said plot except that of a bare licensee and shall not without the previous permission in writing of the (1st Party) transfer his interest in the area allotted to him either in part or whole except for the purpose of raising loans from the House Building Finance Corporation, authorised banks and insurance companies for construction of building thereon.", self.styles['DHAClause']))
        story.append(Spacer(1, 24))
        
        story.append(Paragraph("3.\tShould the licensee commit breach of any of the terms and conditions of these presents or should he neglect to comply with any direction given to him by the 1st Party or in any other respect fail to carry out his obligations under these presents for reasons not beyond his control or fail to pay development charges or other dues, the 1st Party shall have the right to terminate this licence and on such termination the payment made by him to the 1st Party shall be forfeited and he shall have no further claim whatsoever against the 1st Party except at the option of 1st Party he may receive compensation to the extent of the amount of actual expenditure incurred by him on the plot.", self.styles['DHAClause']))
        story.append(Spacer(1, 24))
        
        story.append(Paragraph("Provided that the 1st Party may in his absolute discretion have the building sold out either by public auction or private contract, in which case the licensee shall be entitled to the net sale proceeds of the building or to the amount of actual expenditure incurred by him on having the building constructed whichever is less.", self.styles['DHABody']))
        story.append(Spacer(1, 36))
        
        # Add schedule
        story.append(Paragraph("THE SCHEDULE ABOVE REFERRED TO", self.styles['DHAHeading']))
        story.append(Spacer(1, 24))
        
        schedule = f"""ALL that piece and parcel of land measuring {self.form_data['plot_area']} square yards bearing Plot No {self.form_data['plot_number']} and bounded
North................... {self.form_data['north_boundary']}
South ................{self.form_data['south_boundary']}
East................... {self.form_data['east_boundary']}
West................... {self.form_data['west_boundary']}
Situated in Police Station {self.form_data['police_station']}"""
        
        story.append(Paragraph(schedule, self.styles['DHABody']))
        
        # Add page 4
        story.append(PageBreak())
        story.append(Paragraph("4", self.styles['DHAPageNumber']))
        story.append(Spacer(1, 48))
        
        # Add signature section with proper spacing
        signature_section = """SIGNED by the Secretary, Pakistan\t\t\t\t    …………………………………………...
Defence Officers Housing Authority Karachi.\t\t\t               Signature of the Secretary
In the presence of :\t\t\t\t\t\t                 1st Party"""
        
        story.append(Paragraph(signature_section, self.styles['DHABody']))
        story.append(Spacer(1, 72))
        
        # Add witness section
        witness_section = f"""SIGNED by the above named\t\t\t\t    …………………………………………...
Licensee – 2nd Party in the\t\t\t\t\t        Licensee / 2nd Party
Presence of :

Witness: (1)  Signature:……………………………..
\tName: {self.form_data['witness1_name']}
\tAddress: {self.form_data['witness1_address']}
\tCNIC #: {self.form_data['witness1_cnic']}


  (2)  Signature:……………………………..
\tName: {self.form_data['witness2_name']}
\tAddress: {self.form_data['witness2_address']}
\tCNIC #: {self.form_data['witness2_cnic']}"""
        
        story.append(Paragraph(witness_section, self.styles['DHABody']))
        
        return story
    
    def save(self, output_path: str):
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch
        )
        doc.build(self.generate())

class DHALicenseGenerator:
    def __init__(self, form_data: dict, output_format: OutputFormat = OutputFormat.DOCX):
        self.form_data = form_data
        self.output_format = output_format
        self._generator = self._create_generator()
    
    def _create_generator(self) -> BaseDocumentGenerator:
        if self.output_format == OutputFormat.DOCX:
            return DocxGenerator(self.form_data)
        elif self.output_format == OutputFormat.PDF:
            return PDFGenerator(self.form_data)
        elif self.output_format == OutputFormat.MARKDOWN:
            return MarkdownGenerator(self.form_data)
        elif self.output_format == OutputFormat.HTML:
            return HtmlGenerator(self.form_data)
        else:
            raise ValueError(f"Unsupported output format: {self.output_format}")
    
    def generate(self):
        return self._generator.generate()
    
    def save(self, output_path: str):
        self._generator.save(output_path)

class DocumentGeneratorFactory:
    _generators = {
        FormType.DHA_LICENSE_A: DHALicenseGenerator
    }
    
    @staticmethod
    def create_generator(form_type: FormType, form_data: dict, output_format: str = "docx") -> BaseDocumentGenerator:
        # Validate output format
        try:
            output_format = OutputFormat(output_format.lower())
        except ValueError:
            raise ValueError(f"Invalid output format: {output_format}. Supported formats are: {', '.join([f.value for f in OutputFormat])}")
            
        generator_class = DocumentGeneratorFactory._generators.get(form_type)
        if not generator_class:
            raise ValueError(f"No document generator registered for form type: {form_type}")
        
        # Create generator for the specific format
        return generator_class(form_data, output_format)
    
    @staticmethod
    def register_generator(form_type: FormType, generator_class: type):
        DocumentGeneratorFactory._generators[form_type] = generator_class 